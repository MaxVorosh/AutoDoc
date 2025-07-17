from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_BREAK
from copy import deepcopy

def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None


class Doc:
    def __init__(self, date, voter, target, number, reason, initiator):
        self.date = date
        self.voter = voter
        self.target = target
        self.reason = reason
        self.initiator = initiator
        self.number = number

    def isLong(self):
        maxSymbols = [38, 20, 80]
        currentSymbols = [len(self.voter), len(self.target),
                          len(self.reason) + len(self.initiator) + (self.initiator != '')]
        for i in range(len(maxSymbols)):
            if maxSymbols[i] < currentSymbols[i]:
                return True
        return False


class DocGroup:
    def __init__(self):
        self.docs = []

    def addDoc(self, doc):
        self.docs.append(doc)

    def trimName(self, name):
        prop = ".docx"
        endLen = len(prop)
        while len(name) > endLen and name[-endLen:] == prop:
            name = name[:-endLen]
        name += ".docx"
        return name

    def save(self, name):
        name = self.trimName(name)
        self.saveTemplate(name)
        d = Document(name)
        for ti in range(len(d.tables)):
            t = d.tables[ti]
            width = len(t.columns)
            height = len(t.rows)
            for i in range(width):
                for j in range(height):
                    cell = t.cell(i, j)
                    self.updateCell(cell, self.docs[ti])
        d.save(name)

    def updateCell(self, cell, info):
        if cell.text == "<date>":
            cell.text = info.date
        elif cell.text == "<voter>":
            cell.text = info.voter
        elif cell.text == "<reason>":
            cell.text = f"{info.reason} {info.initiator}"
        elif "<target>" in cell.text:
            cell.text = cell.text.replace("<target>", info.target).replace("<number>", info.number)
        else:
            return
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        cell.paragraphs[0].runs[0].font.name = 'Times New Roman'

    def saveTemplate(self, name):
        d = Document("data/baseDoc.docx")
        second = False
        for i in range(len(self.docs) - 1):
            second = not second
            newPage = False
            p = d.paragraphs[-1]
            if self.docs[i].isLong() or self.docs[i + 1].isLong() and second:
                d.add_page_break()
                second = False
                newPage = True
            if second and not newPage:
                r = p.add_run("-" * 100)
                r.font.size = Pt(14)
                r.font.color.rgb = RGBColor(200, 200, 200)
            p = d.paragraphs[-1]
            p._p.addnext(deepcopy(d.tables[0]._tbl))
            if not second or newPage:
                d.add_paragraph()
        d.save(name)